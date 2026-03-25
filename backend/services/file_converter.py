"""
services/file_converter.py

Converts uploaded files to clean markdown text for BM25 indexing.
Supported: .md, .txt, .pdf, .docx, .csv
"""
import io
import csv


def convert_to_markdown(content: bytes, filename: str) -> str:
    """
    Convert file bytes to a clean markdown string.
    Returns the markdown text or raises ValueError on unsupported type.
    """
    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''

    if ext in ('md', 'txt'):
        return _convert_text(content)

    if ext == 'pdf':
        return _convert_pdf(content, filename)

    if ext == 'docx':
        return _convert_docx(content, filename)

    if ext == 'csv':
        return _convert_csv(content, filename)

    raise ValueError(f"Unsupported file type: .{ext}. Allowed: .md, .txt, .pdf, .docx, .csv")


# ── Text / Markdown ───────────────────────────────────────────────
def _convert_text(content: bytes) -> str:
    try:
        return content.decode('utf-8')
    except UnicodeDecodeError:
        return content.decode('latin-1')


# ── PDF ───────────────────────────────────────────────────────────
def _convert_pdf(content: bytes, filename: str) -> str:
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("pdfplumber is required for PDF support. Add it to requirements.txt")

    pages = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for i, page in enumerate(pdf.pages, 1):
            text = page.extract_text()
            if text and text.strip():
                pages.append(f"## Page {i}\n\n{text.strip()}")

    if not pages:
        raise ValueError(f"Could not extract text from {filename}. The PDF may be scanned/image-based.")

    return "\n\n".join(pages)


# ── DOCX ──────────────────────────────────────────────────────────
def _convert_docx(content: bytes, filename: str) -> str:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required for DOCX support. Add it to requirements.txt")

    doc     = Document(io.BytesIO(content))
    lines   = []

    for para in doc.paragraphs:
        text  = para.text.strip()
        style = para.style.name if para.style else ''

        if not text:
            continue

        if 'Heading 1' in style:
            lines.append(f"# {text}")
        elif 'Heading 2' in style:
            lines.append(f"## {text}")
        elif 'Heading 3' in style:
            lines.append(f"### {text}")
        elif 'List' in style:
            lines.append(f"- {text}")
        else:
            lines.append(text)

    # Also extract tables
    for table in doc.tables:
        rows = []
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            rows.append('| ' + ' | '.join(cells) + ' |')
            if i == 0:
                rows.append('| ' + ' | '.join(['---'] * len(cells)) + ' |')
        if rows:
            lines.append('\n'.join(rows))

    if not lines:
        raise ValueError(f"No readable text found in {filename}")

    return '\n\n'.join(lines)


# ── CSV ───────────────────────────────────────────────────────────
def _convert_csv(content: bytes, filename: str) -> str:
    try:
        text = content.decode('utf-8')
    except UnicodeDecodeError:
        text = content.decode('latin-1')

    reader  = csv.reader(io.StringIO(text))
    rows    = list(reader)

    if not rows:
        raise ValueError(f"No data found in {filename}")

    lines = []
    headers = rows[0] if rows else []

    # Build markdown table
    if headers:
        lines.append('| ' + ' | '.join(headers) + ' |')
        lines.append('| ' + ' | '.join(['---'] * len(headers)) + ' |')
        for row in rows[1:]:
            # Pad row if shorter than headers
            padded = row + [''] * (len(headers) - len(row))
            lines.append('| ' + ' | '.join(padded[:len(headers)]) + ' |')

    if not lines:
        raise ValueError(f"Could not parse {filename}")

    return f"# {filename}\n\n" + '\n'.join(lines)
